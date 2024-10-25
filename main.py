
import gradio as gr
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from openai import OpenAI
import os
import time
from selenium.webdriver.common.by import By
from selenium.common.exceptions import WebDriverException
from docx import Document
from docx.shared import Pt
from diffusers import FluxPipeline
from diffusers import StableDiffusionPipeline
import torch
from image import *
import base64


# Установка API ключа
os.environ["OPENAI_API_KEY"] = ""

# Загрузка модели
client = OpenAI(
    base_url="https://api.sambanova.ai/v1/",
    api_key=os.getenv("OPENAI_API_KEY"),
)
model = "Meta-Llama-3.1-405B-Instruct"

# Запуск chromedriver
chrome_options = Options()
chrome_options.add_argument("--headless")  # Запуск без интерфейса (по желанию)
chrome_options.add_argument("--ignore-certificate-errors")
service = Service("C:/chromedriver/chromedriver.exe") 

driver = webdriver.Chrome(service=service, options=chrome_options)

# 1. Функция поиска информации в Google за последний месяц по запросу query
def search_google(query):
    driver.get(f"https://www.google.com/search?q={query}&tbs=qdr:m")
    time.sleep(5)
    results = driver.find_elements(By.CSS_SELECTOR, 'h3')
    links = []
    for result in results[-5:]:
        link = result.find_element(By.XPATH, '..').get_attribute('href')  # Получаем ссылку
        links.append(link)
    page_content = []

    for link in links:
        try:
            driver.get(link)
            time.sleep(2)
            text = driver.find_element(By.TAG_NAME, 'body').text
            page_content.append(text[:500])
        except WebDriverException as e:
            print(f"Ошибка при открытии {link}: {e}")
    return "\n".join(page_content)

# 2. Генерация текста для блога по найденной информации
def gen(prompt):
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "Ты помощник, у которого целевая аудитория - мамы"},
            {"role": "user", "content": prompt}
        ],
        stream=True,
        temperature=0.5,
    )
    response = ""
    for chunk in completion:
        response += chunk.choices[0].delta.content or ""
    return response

# Запуск основного процесса поиски информации и её геенрации
def run_process():
    query = "Блог о важности раннего обучения детей арифметике, скорочтению и другим предметам"
    answer = search_google(query)
    prompt_1 = f"На основе следующей информации об обучении детей с раннего возраста арифметике, скорочтению и другим предметам: {answer} \n создайте краткий и содержательный текст для блога."
    text_response = gen(prompt_1)
    return text_response

# 3. Получение основной мысли по найденному тексту
def extract_main_idea(text_response):
    prompt_image = f"Выдели краткую мысль (не более 10 слов) из текста на русском языке {text_response}."
    main_idea = gen(prompt_image)
    return main_idea

# 4. Генерация изображения по основной мысли текста
def generate_image(main_idea):
    api = Text2ImageAPI('https://api-key.fusionbrain.ai/', '', '')
    model_id = api.get_model()
    uuid = api.generate(main_idea, model_id)
    images = api.check_generation(uuid)
    # Здесь image_base64 - это строка с данными изображения в формате base64
    image_base64 = images[0] 
    # Декодируем строку base64 в бинарные данные
    image_data = base64.b64decode(image_base64)
    # Открываем файл для записи бинарных данных изображения
    with open("image.jpg", "wb") as file:
        file.write(image_data)
    return "image.jpg"


# 5. Сохранение данных
def save_data(text_response, main_idea):
    doc = Document()
    doc.add_heading(f'{main_idea}', level=1)
    run = doc.add_paragraph().add_run(f'{text_response}')
    run.font.size = Pt(12)
    run.bold = True
    run.italic = True
    doc.save('Блог.docx')
    with open('текст-для-изображения.txt', 'w', encoding='utf-8') as file:
        file.write(main_idea)
    return "Данные успешно сохранены!"

# Настройка интерфейса Gradio
with gr.Blocks() as iface:
    run_button = gr.Button("Запуск")
    text_output = gr.Textbox(label="Сгенерированный текст", interactive=False)
    run_button.click(run_process, outputs=text_output)

    extract_button = gr.Button("Основная мысль")
    main_idea_output = gr.Textbox(label="Основная мысль", interactive=False)
    extract_button.click(extract_main_idea, inputs=text_output, outputs=main_idea_output)
    
    generate_button = gr.Button("Генерация изображения")
    image_output = gr.Image(label="Сгенерированное изображение", interactive=False)
    generate_button.click(generate_image, inputs=main_idea_output, outputs=image_output)
    
    save_button = gr.Button("Сохранить")
    save_output = gr.Textbox(label="Сохранено", interactive=False)
    save_button.click(save_data,inputs=[text_output,main_idea_output ], outputs=save_output)

if __name__ == "__main__":
    iface.launch()

# Закрытие драйвера
driver.quit()
    

